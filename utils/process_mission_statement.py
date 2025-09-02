from copy import deepcopy
from datetime import date
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Pt, RGBColor, Cm
import io
import pandas as pd
from pathlib import Path
from PIL import Image
import streamlit as st


def statement_checklist():
    # Section 1: Request and notifications
    checklist_1 = {"Approved mission request": (True if "mission_request" in st.session_state else False),
                   "F4E Mission notification": (True if "notification" in st.session_state else False)}
    if st.session_state.mission_modified and not "modified_mission_request" in st.session_state:
        checklist_1.update({"Modified mission request": False})
    elif st.session_state.mission_modified and "modified_mission_request" in st.session_state:
        checklist_1.update({"Modified mission request": True})
    if st.session_state.mission_cancelled and not "cancellation" in st.session_state:
        checklist_1.update({"Mission cancellation notification": False})
    elif st.session_state.mission_cancelled and "cancellation" in st.session_state:
        checklist_1.update({"Mission cancellation notification": True})
    
    # Section 2: Expenses
    checklist_2 = {}
    mission_data = deepcopy(st.session_state.mission_data)
    if not mission_data is None:
        # Accommodation
        if mission_data["number of days"] > 1 and "Accommodation" in list(st.session_state.expenses["Category"]):
            checklist_2.update({"Accommodation": True})
        elif mission_data["number of days"] > 1 and not "Accommodation" in list(st.session_state.expenses["Category"]):
            checklist_2.update({"Accommodation": False})
        # Travel

        
    return


def generate_mission_statement():
    filename_out = st.session_state.filename_out
    mission_data = st.session_state.mission_data
    title_label, employee_name = [mission_data[key] for key in ["title label", "employee name"]]
    mission_request = deepcopy(st.session_state.mission_request)
    if "modified_mission_request" in st.session_state:
        modified_mission_request = deepcopy(st.session_state.modified_mission_request)
    else:
        modified_mission_request = None
    notification = deepcopy(st.session_state.notification)
    img = Image.open(fp=notification)
    img_size_pix = img.size
    del img
    if img_size_pix[0] / img_size_pix[1] > 0.7071:
        notification_size = (Cm(15.0), None)
    else:
        notification_size = (None, Cm(21.0))
    if "cancellation" in st.session_state:
        cancellation = deepcopy(st.session_state.cancellation)
        img = Image.open(fp=cancellation)
        img_size_pix = img.size
        del img
        if img_size_pix[0] / img_size_pix[1] > 0.7071:
            cancellation_size = (Cm(15.0), None)
        else:
            cancellation_size = (None, Cm(21.0))
    else:
        cancellation = None
    expenses_df = deepcopy(st.session_state.expenses)
    summary_df = deepcopy(st.session_state.expense_summary)
    attachments_dict = deepcopy(st.session_state.attachments)
    kilometers_own_car = st.session_state.kilometers_own_car

    collective = False
    currency_flag = False

    doc = Document()
    norm_style = doc.styles["Normal"]
    norm_style.font.name = "Calibri"
    norm_style.font.size = Pt(11)
    norm_style.font.color.rgb = RGBColor(0, 0, 0)
    norm_style.font.bold = False
    norm_style.font.underline = False
    doc.sections[0].footer.add_paragraph("ATG Public")
    doc.sections[0].footer.paragraphs[-1].runs[0].font.name = "Calibri"
    doc.sections[0].footer.paragraphs[-1].runs[0].font.size = Pt(10)
    doc.sections[0].footer.paragraphs[-1].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.sections[0].footer.paragraphs[-1].paragraph_format.left_indent = -Cm(1.75)
    doc.sections[0].footer_distance = Cm(0.25)
    doc.sections[0].top_margin = Cm(2.54)
    doc.sections[0].right_margin = Cm(2.54)
    doc.sections[0].bottom_margin = Cm(2.54)
    doc.sections[0].left_margin = Cm(2.54)

    # Title
    title = doc.add_paragraph(title_label)
    title.runs[0].font.size = Pt(14)
    title.runs[0].font.bold = True
    title.runs[0].font.underline = True
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Mission request section
    tbl_1 = doc.add_table(rows=1, cols=1)
    tbl_1.style = "Table Grid"
    req_cell = tbl_1.cell(0, 0)
    req_cell.paragraphs[0].text = "Screenshot of the approved Mission Request"
    req_cell.paragraphs[0].runs[0].font.bold = True
    for img in mission_request:
        req_cell.add_paragraph().add_run().add_picture(img, height=Cm(21.0))
        req_cell.paragraphs[-1].add_run(" ")
    req_cell.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if modified_mission_request:
        req_cell.add_paragraph("Screenshot of the modified Mission Request")
        req_cell.paragraphs[-1].runs[0].font.bold = True
        for img in modified_mission_request:
            req_cell.add_paragraph().add_run().add_picture(img, height=Cm(21.0))
            req_cell.paragraphs[-1].add_run(" ")
        req_cell.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    req_cell.add_paragraph("Screenshot of the notification/greenlight from F4E to go on the mission")
    req_cell.paragraphs[-1].runs[0].font.bold = True
    if "notification_comment" in st.session_state:
        req_cell.add_paragraph("NOTE: " + st.session_state.notification_comment)
    req_cell.add_paragraph().add_run().add_picture(notification, width=notification_size[0], height=notification_size[1])
    req_cell.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if cancellation:
        req_cell.add_paragraph("Screenshot of the cancellation notification")
        req_cell.paragraphs[-1].runs[0].font.bold = True
        req_cell.add_paragraph().add_run().add_picture(cancellation, width=cancellation_size[0], height=cancellation_size[1])
        req_cell.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


    # Expenses section
    additional_label = doc.add_paragraph("Additional information:")
    additional_label.runs[0].font.bold = True
    additional_label.runs[0].font.underline = True
    additional_label.paragraph_format.page_break_before = True

    tbl_2 = doc.add_table(rows=1, cols=1)
    tbl_2.style = "Table Grid"
    expense_cell = tbl_2.cell(0, 0)
    expense_cell.paragraphs[0].text = "Screenshots of boarding passes, meal tickets, hotels, train ticket, etc"
    expense_cell.paragraphs[0].runs[0].font.bold = True
    # Accommodation
    total = float(summary_df[summary_df["Category"] == "Accommodation"].Amount.iloc[0])
    if total > 0.0:
        expense_cell.add_paragraph("Accommodation: {:.2f} ".format(total) + u"\N{euro sign}")
    else:
        expense_cell.add_paragraph("Accommodation")
    expense_cell.paragraphs[-1].runs[0].font.bold = True
    idx_accommodation = expenses_df.index[expenses_df["Category"] == "Accommodation"].to_list()
    if idx_accommodation:
        expense_cell.add_paragraph()
        expense_cell.paragraphs[-1].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for idx in idx_accommodation:
            if "Exchange rate" in expenses_df:
                if expenses_df.loc[idx, "Currency"] != "EUR":
                    expense_cell.add_paragraph("{}: {:.2f} {} x {} EUR/{} = {:.2f} EUR".format(expenses_df.loc[idx, "Description"], 
                                                                                                expenses_df.loc[idx, "Paid amount"],
                                                                                                expenses_df.loc[idx, "Currency"],
                                                                                                expenses_df.loc[idx, "Exchange rate"],
                                                                                                expenses_df.loc[idx, "Currency"],
                                                                                                expenses_df.loc[idx, "Paid amount (EUR)"]))
                    expense_cell.paragraphs[-1].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    expense_cell.add_paragraph()
                    expense_cell.paragraphs[-1].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    currency_flag = True
                else:
                    currency_flag = False
            for img, size in zip(*attachments_dict[idx]):
                width = (Cm(size[0]) if size[0] else None)
                height = (Cm(size[1]) if size[1] else None)
                expense_cell.paragraphs[-1].add_run().add_picture(img, width=width, height=height)
                if currency_flag:
                    expense_cell.add_paragraph()
                    expense_cell.paragraphs[-1].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    expense_cell.paragraphs[-1].add_run(" ")
        empty_section = False
    else:
        expense_cell.add_paragraph("N/A")
        expense_cell.paragraphs[-1].runs[0].font.bold = True
        expense_cell.paragraphs[-1].runs[0].font.color.rgb = RGBColor(255, 0, 0)
        empty_section = True

    # Travel
    total = float(summary_df[summary_df["Category"] == "Travel (flight, train, taxi, ...)"].Amount.iloc[0])
    # if empty_section:
    expense_cell.add_paragraph()
    if total > 0.0:
        expense_cell.add_paragraph("Travel: {:.2f} ".format(total) + u"\N{euro sign}")
    else:
        expense_cell.add_paragraph("Travel")
    # else:
    #     doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    #     tbl_3 = doc.add_table(rows=1, cols=1)
    #     tbl_3.style = "Table Grid"
    #     expense_cell = tbl_3.cell(0, 0)
    #     if total > 0.0:
    #         expense_cell.paragraphs[0].text = "Travel: {:.2f}".format(total) + u"\N{euro sign}"
    #     else:
    #         expense_cell.paragraphs[0].text = "Travel"
    expense_cell.paragraphs[-1].runs[0].font.bold = True
    idx_travel = expenses_df.index[(expenses_df["Category"] == "Flights") |
                                    (expenses_df["Category"] == "Train") |
                                    (expenses_df["Category"] == "Bus") |
                                    (expenses_df["Category"] == "Metro") |
                                    (expenses_df["Category"] == "Taxi") |
                                    (expenses_df["Category"] == "Rental car") |
                                    (expenses_df["Category"] == "Own car") |
                                    (expenses_df["Category"] == "Additional travel expenses")].to_list()
    idx_own_car = expenses_df.index[expenses_df["Category"] == "Own car"].to_list()
    if idx_travel:
        expense_cell.add_paragraph()
        expense_cell.paragraphs[-1].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for idx in idx_travel:
            if "Exchange rate" in expenses_df:
                if expenses_df.loc[idx, "Currency"] != "EUR" and expenses_df.loc[idx, "Category"] != "Own car":
                    expense_cell.add_paragraph("{}: {:.2f} {} x {} EUR/{} = {:.2f} EUR".format(expenses_df.loc[idx, "Description"], 
                                                                                                expenses_df.loc[idx, "Paid amount"],
                                                                                                expenses_df.loc[idx, "Currency"],
                                                                                                expenses_df.loc[idx, "Exchange rate"],
                                                                                                expenses_df.loc[idx, "Currency"],
                                                                                                expenses_df.loc[idx, "Paid amount (EUR)"]))
                    expense_cell.paragraphs[-1].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    expense_cell.add_paragraph()
                    expense_cell.paragraphs[-1].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    currency_flag = True
                else:
                    currency_flag = False
            if idx in idx_own_car:
                expense_cell.add_paragraph("Own car: " + expenses_df.loc[idx]["Description"])
                kms = float(kilometers_own_car[idx])
                if expenses_df.shape[1] == 7:
                    euros = float(expenses_df.loc[idx]["Paid amount (EUR)"])
                else:
                    euros = float(expenses_df.loc[idx]["Paid amount"])
                expense_cell.add_paragraph("Kilometer calculation: {:.1f} km x 0.28 ".format(kms) + u"\N{euro sign}" + "/km = {:.2f} ".format(euros) + u"\N{euro sign}")
                expense_cell.paragraphs[-1].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                expense_cell.add_paragraph()
                expense_cell.paragraphs[-1].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for img, size in zip(*attachments_dict[idx]):
                width = (Cm(size[0]) if size[0] else None)
                height = (Cm(size[1]) if size[1] else None)
                expense_cell.paragraphs[-1].add_run().add_picture(img, width=width, height=height)
                if currency_flag:
                    expense_cell.add_paragraph()
                    expense_cell.paragraphs[-1].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    expense_cell.paragraphs[-1].add_run(" ")
        empty_section = False
    else:
        expense_cell.add_paragraph("N/A")
        expense_cell.paragraphs[-1].runs[0].font.bold = True
        expense_cell.paragraphs[-1].runs[0].font.color.rgb = RGBColor(255, 0, 0)
        empty_section = True

    # Meals
    total = float(summary_df[summary_df["Category"] == "Meals"].Amount.iloc[0])
    # if empty_section:
    # expense_cell.add_paragraph()
    # if total > 0.0:
    #     expense_cell.add_paragraph("Meals: {:.2f}".format(total) + u"\N{euro sign}")
    # else:
    #     expense_cell.add_paragraph("Meals")
    # else:
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    tbl_4 = doc.add_table(rows=1, cols=1)
    tbl_4.style = "Table Grid"
    expense_cell = tbl_4.cell(0, 0)
    if total > 0.0:
        expense_cell.paragraphs[0].text = "Meals: {:.2f} ".format(total) + u"\N{euro sign}"
    else:
        expense_cell.paragraphs[0].text = "Meals"
    expense_cell.paragraphs[-1].runs[0].font.bold = True
    idx_meals = expenses_df.index[expenses_df["Category"] == "Meals"].to_list()
    if idx_meals:
        expense_cell.add_paragraph()
        expense_cell.paragraphs[-1].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for idx in idx_meals:
            if "Exchange rate" in expenses_df:
                if expenses_df.loc[idx, "Currency"] != "EUR":
                    expense_cell.add_paragraph("{}: {:.2f} {} x {} EUR/{} = {:.2f} EUR".format(expenses_df.loc[idx, "Description"], 
                                                                                                expenses_df.loc[idx, "Paid amount"],
                                                                                                expenses_df.loc[idx, "Currency"],
                                                                                                expenses_df.loc[idx, "Exchange rate"],
                                                                                                expenses_df.loc[idx, "Currency"],
                                                                                                expenses_df.loc[idx, "Paid amount (EUR)"]))
                    expense_cell.paragraphs[-1].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    expense_cell.add_paragraph()
                    expense_cell.paragraphs[-1].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    currency_flag = True
                else:
                    currency_flag = False
            for img, size in zip(*attachments_dict[idx]):
                width = (Cm(size[0]) if size[0] else None)
                height = (Cm(size[1]) if size[1] else None)
                expense_cell.paragraphs[-1].add_run().add_picture(img, width=width, height=height)
                if currency_flag:
                    expense_cell.add_paragraph()
                    expense_cell.paragraphs[-1].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    expense_cell.paragraphs[-1].add_run(" ")
    else:
        expense_cell.add_paragraph("N/A")
        expense_cell.paragraphs[-1].runs[0].font.bold = True
        expense_cell.paragraphs[-1].runs[0].font.color.rgb = RGBColor(255, 0, 0)

    # Expense summary
    expense_cell.add_paragraph()
    expense_cell.add_paragraph()
    expense_summary = expense_cell.add_table(rows=5, cols=3)
    expense_summary.style = "Table Grid"
    widths = [Cm(1.08), Cm(8.12), Cm(6.9)]
    for row in expense_summary.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width
    expense_summary.cell(0, 0).paragraphs[0].text = "No."
    expense_summary.cell(0, 0).paragraphs[0].runs[0].font.bold = True
    expense_summary.cell(0, 1).paragraphs[0].text = "Item"
    expense_summary.cell(0, 1).paragraphs[0].runs[0].font.bold = True
    expense_summary.cell(0, 2).paragraphs[0].text = "Total cost (" + u"\N{euro sign}" + ")"
    expense_summary.cell(0, 2).paragraphs[0].runs[0].font.bold = True
    for idx, cell in enumerate(expense_summary.columns[0].cells[1:-1]):
        cell.paragraphs[0].text = str(idx + 1)
    item_labels = ["Accommodation", "Travel (Car rental, flights, train, fuel, tolls, etc.)", "Meals"]
    for cell, label in zip(expense_summary.columns[1].cells[1:-1], item_labels):
        cell.paragraphs[0].text = label
    total_cell = expense_summary.cell(4, 0).merge(expense_summary.cell(4, 1))
    total_cell.paragraphs[0].text = "Total"
    total_cell.paragraphs[0].runs[0].font.bold = True
    total_cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    expense_summary.rows[4].height = Cm(0.93)
    expense_summary.cell(1, 2).paragraphs[0].text = "{:.2f}".format(float(summary_df[summary_df["Category"] == "Accommodation"].Amount.iloc[0]))
    expense_summary.cell(1, 2).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    expense_summary.cell(1, 2).paragraphs[0].paragraph_format.right_indent = Cm(5.0)    
    expense_summary.cell(2, 2).paragraphs[0].text = "{:.2f}".format(float(summary_df[summary_df["Category"] == "Travel (flight, train, taxi, ...)"].Amount.iloc[0]))
    expense_summary.cell(2, 2).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    expense_summary.cell(2, 2).paragraphs[0].paragraph_format.right_indent = Cm(5.0)    
    expense_summary.cell(3, 2).paragraphs[0].text = "{:.2f}".format(float(summary_df[summary_df["Category"] == "Meals"].Amount.iloc[0]))
    expense_summary.cell(3, 2).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    expense_summary.cell(3, 2).paragraphs[0].paragraph_format.right_indent = Cm(5.0)    
    expense_summary.cell(4, 2).paragraphs[0].text = "{:.2f}".format(float(summary_df[summary_df["Category"] == "Total"].Amount.iloc[0]))
    expense_summary.cell(4, 2).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    expense_summary.cell(4, 2).paragraphs[0].paragraph_format.right_indent = Cm(5.0)    

    
    expense_cell.paragraphs[-1].text = "*Summary of expenses is not to be understood as fully " \
                                       "comprehensive as certain costs are not reported in this Mission Statement"
    expense_cell.paragraphs[-1].runs[0].font.italic = True
    if collective:
        expense_cell.add_paragraph("NOTE: This is a collective mission with F4E-OMF-1159-01-01-XX Name Mission Statement # ")
        expense_cell.paragraphs[-1].runs[0].font.bold = True
        expense_cell.paragraphs[-1].runs[0].font.color.rgb = RGBColor(255, 0, 0)


    # Signatures
    signatures_label = doc.add_paragraph("Signatures:")
    signatures_label.runs[0].font.bold = True
    signatures_label.runs[0].font.underline = True
    signatures_label.paragraph_format.page_break_before = False
    signatures_table = doc.add_table(rows=3, cols=2)
    signatures_table.style = "Table Grid"
    widths = [Cm(6.97), Cm(9.57)]
    for row in signatures_table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width
    signatures_table.rows[0].height = Cm(1.24)
    signatures_table.rows[1].height = Cm(1.22)
    signatures_table.rows[2].height = Cm(0.51)
    signatures_table.cell(0, 0).paragraphs[0].text = "Participant/s:"
    signatures_table.cell(0, 0).paragraphs[0].runs[0].bold = True
    signatures_table.cell(0, 0).add_paragraph(employee_name)
    signatures_table.cell(1, 0).paragraphs[0].text = "Name and Signature:"
    signatures_table.cell(1, 0).paragraphs[0].runs[0].bold = True
    signatures_table.cell(1, 0).add_paragraph("PjM/ Rose Marcaida")
    signatures_table.cell(2, 0).paragraphs[0].text = "Date: {}".format(date.today().strftime('%d/%m/%Y'))
    signatures_table.cell(2, 0).paragraphs[0].runs[0].bold = True
    signatures_table.cell(0, 1).paragraphs[0].text = "F4E:"
    signatures_table.cell(0, 1).paragraphs[0].runs[0].bold = True
    signatures_table.cell(1, 1).paragraphs[0].text = "Name and Signature:"
    signatures_table.cell(1, 1).paragraphs[0].runs[0].bold = True
    signatures_table.cell(2, 1).paragraphs[0].text = "Date:"
    signatures_table.cell(2, 1).paragraphs[0].runs[0].bold = True


    f = io.BytesIO()
    f.name = filename_out
    doc.save(f)
    return f
